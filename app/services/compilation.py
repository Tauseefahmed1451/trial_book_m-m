"""Final draft compilation service."""

from __future__ import annotations

import io
import uuid

from docx import Document
from sqlalchemy.orm import Session

from app.db import models
from app.schemas.books import ArtifactResponse
from app.services.storage import LocalStorageProvider


class CompilationService:
    """Compile approved chapters into output artifacts."""

    def __init__(self, db: Session) -> None:
        self.db = db
        self.storage = LocalStorageProvider()

    def compile_book(self, book_id: uuid.UUID) -> list[ArtifactResponse]:
        """Compile the latest chapter versions into txt and docx files."""
        book = self.db.get(models.Book, book_id)
        if book is None:
            raise ValueError("Book not found")

        latest_versions = []
        for chapter in sorted(book.chapters, key=lambda item: item.chapter_number):
            if not chapter.versions:
                continue
            latest_versions.append(max(chapter.versions, key=lambda item: item.version_no))

        if not latest_versions:
            raise ValueError("No chapters available for compilation")

        full_text_parts = [book.title, ""]
        for chapter, version in zip(sorted(book.chapters, key=lambda item: item.chapter_number), latest_versions, strict=False):
            full_text_parts.append(f"Chapter {chapter.chapter_number}: {chapter.chapter_title}")
            full_text_parts.append(version.content_text)
            full_text_parts.append("")
        full_text = "\n".join(full_text_parts)

        txt_path = self.storage.save_bytes(
            book_id=str(book.id),
            file_name="final_draft.txt",
            content=full_text.encode("utf-8"),
        )
        txt_artifact = models.BookArtifact(
            book_id=book.id,
            artifact_type="txt",
            storage_backend="local",
            storage_path=txt_path,
            file_name="final_draft.txt",
        )
        self.db.add(txt_artifact)

        document = Document()
        document.add_heading(book.title, level=1)
        for chapter, version in zip(sorted(book.chapters, key=lambda item: item.chapter_number), latest_versions, strict=False):
            document.add_heading(f"Chapter {chapter.chapter_number}: {chapter.chapter_title}", level=2)
            document.add_paragraph(version.content_text)
        buffer = io.BytesIO()
        document.save(buffer)
        docx_path = self.storage.save_bytes(
            book_id=str(book.id),
            file_name="final_draft.docx",
            content=buffer.getvalue(),
        )
        docx_artifact = models.BookArtifact(
            book_id=book.id,
            artifact_type="docx",
            storage_backend="local",
            storage_path=docx_path,
            file_name="final_draft.docx",
        )
        self.db.add(docx_artifact)

        book.current_stage = models.BookStage.COMPILED
        book.current_status = models.BookStatus.COMPLETED
        self.db.flush()
        return [
            ArtifactResponse.model_validate(txt_artifact, from_attributes=True),
            ArtifactResponse.model_validate(docx_artifact, from_attributes=True),
        ]
